import uuid
import json
from io import BytesIO
import base64
import boto3
from docx import Document
from PIL import Image, UnidentifiedImageError
from botocore.config import Config as BotoConfig
from botocore.exceptions import ClientError
from app.config import settings

CHAT_ASSET_ALLOWED_MIME_TYPES = {
    "image/png": "png",
    "image/jpeg": "jpg",
    "image/webp": "webp",
    "image/gif": "gif",
}
CHAT_ASSET_PREVIEW_MAX_SIZE = (1440, 1440)
CHAT_ASSET_PREVIEW_QUALITY = 82


def _make_s3_client():
    return boto3.client(
        "s3",
        endpoint_url=f"{'https' if settings.minio_use_ssl else 'http'}://{settings.minio_endpoint}",
        aws_access_key_id=settings.minio_access_key,
        aws_secret_access_key=settings.minio_secret_key,
        config=BotoConfig(signature_version="s3v4"),
        region_name="us-east-1",
    )


s3 = _make_s3_client()
BUCKET = settings.minio_bucket


def ensure_bucket() -> None:
    try:
        s3.head_bucket(Bucket=BUCKET)
    except ClientError as e:
        error_code = e.response["Error"]["Code"]
        if error_code in ("NoSuchBucket", "404"):
            s3.create_bucket(Bucket=BUCKET)
        else:
            raise


def generate_id() -> str:
    return uuid.uuid4().hex[:12]


def upload_document(filename: str, file_data: bytes) -> str:
    document_id = generate_id()
    s3.put_object(Bucket=BUCKET, Key=f"documents/{document_id}/original.docx", Body=file_data)
    s3.put_object(Bucket=BUCKET, Key=f"documents/{document_id}/current.docx", Body=file_data)
    meta = json.dumps({"name": filename, "document_id": document_id})
    s3.put_object(Bucket=BUCKET, Key=f"documents/{document_id}/meta.json", Body=meta.encode())
    return document_id


def _build_blank_docx(name: str) -> bytes:
    document = Document()
    document.add_heading(name or "Untitled", level=1)
    document.add_paragraph("")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def create_blank_document(name: str) -> str:
    document_id = generate_id()
    blank_docx = _build_blank_docx(name)
    s3.put_object(Bucket=BUCKET, Key=f"documents/{document_id}/original.docx", Body=blank_docx)
    s3.put_object(Bucket=BUCKET, Key=f"documents/{document_id}/current.docx", Body=blank_docx)
    meta = json.dumps({"name": name, "document_id": document_id})
    s3.put_object(Bucket=BUCKET, Key=f"documents/{document_id}/meta.json", Body=meta.encode())
    return document_id


def get_document_info(document_id: str) -> dict:
    resp = s3.get_object(Bucket=BUCKET, Key=f"documents/{document_id}/meta.json")
    return json.loads(resp["Body"].read())


def download_document(document_id: str) -> bytes:
    resp = s3.get_object(Bucket=BUCKET, Key=f"documents/{document_id}/current.docx")
    return resp["Body"].read()


def save_current_document(document_id: str, file_data: bytes) -> None:
    s3.put_object(Bucket=BUCKET, Key=f"documents/{document_id}/current.docx", Body=file_data)


def delete_document(document_id: str) -> bool:
    prefix = f"documents/{document_id}/"
    response = s3.list_objects_v2(Bucket=BUCKET, Prefix=prefix)
    if "Contents" not in response or not response["Contents"]:
        return False
    objects = [{"Key": obj["Key"]} for obj in response["Contents"]]
    s3.delete_objects(Bucket=BUCKET, Delete={"Objects": objects})
    return True


def upload_chat_asset(document_id: str, filename: str, file_data: bytes, mime_type: str) -> dict:
    extension = CHAT_ASSET_ALLOWED_MIME_TYPES.get(mime_type)
    if not extension:
        raise ValueError("Unsupported image type")

    try:
        image = Image.open(BytesIO(file_data))
        image.load()
    except UnidentifiedImageError as exc:
        raise ValueError("Unsupported or invalid image content") from exc
    width, height = image.size

    asset_id = generate_id()
    asset_prefix = _chat_asset_prefix(document_id, asset_id)
    original_key = f"{asset_prefix}/original.{extension}"
    preview_key = f"{asset_prefix}/preview.jpg"
    meta_key = f"{asset_prefix}/meta.json"

    preview_bytes = _build_chat_asset_preview(image)
    metadata = {
        "asset_id": asset_id,
        "filename": filename,
        "mime_type": mime_type,
        "width": width,
        "height": height,
        "size_bytes": len(file_data),
        "storage_key": original_key,
        "preview_storage_key": preview_key,
    }

    s3.put_object(Bucket=BUCKET, Key=original_key, Body=file_data, ContentType=mime_type)
    s3.put_object(Bucket=BUCKET, Key=preview_key, Body=preview_bytes, ContentType="image/jpeg")
    s3.put_object(Bucket=BUCKET, Key=meta_key, Body=json.dumps(metadata).encode("utf-8"), ContentType="application/json")
    return metadata


def get_chat_asset_info(document_id: str, asset_id: str) -> dict:
    response = s3.get_object(Bucket=BUCKET, Key=f"{_chat_asset_prefix(document_id, asset_id)}/meta.json")
    return json.loads(response["Body"].read())


def get_chat_asset_bytes(document_id: str, asset_id: str, variant: str = "original") -> bytes:
    metadata = get_chat_asset_info(document_id, asset_id)
    if variant == "preview":
        key = metadata["preview_storage_key"]
    else:
        key = metadata["storage_key"]

    response = s3.get_object(Bucket=BUCKET, Key=key)
    return response["Body"].read()


def get_chat_asset_preview_data_url(document_id: str, asset_id: str) -> str:
    preview_bytes = get_chat_asset_bytes(document_id, asset_id, variant="preview")
    encoded = base64.b64encode(preview_bytes).decode("ascii")
    return f"data:image/jpeg;base64,{encoded}"


def delete_chat_asset(document_id: str, asset_id: str) -> bool:
    prefix = f"{_chat_asset_prefix(document_id, asset_id)}/"
    response = s3.list_objects_v2(Bucket=BUCKET, Prefix=prefix)
    if "Contents" not in response or not response["Contents"]:
        return False

    objects = [{"Key": obj["Key"]} for obj in response["Contents"]]
    s3.delete_objects(Bucket=BUCKET, Delete={"Objects": objects})
    return True


def _chat_asset_prefix(document_id: str, asset_id: str) -> str:
    return f"documents/{document_id}/chat-assets/{asset_id}"


def _build_chat_asset_preview(image: Image.Image) -> bytes:
    preview = image.copy()
    if getattr(preview, "is_animated", False):
        preview.seek(0)
        preview = preview.copy()

    if preview.mode not in ("RGB", "L"):
        rgba_preview = preview.convert("RGBA")
        background = Image.new("RGB", preview.size, "white")
        background.paste(rgba_preview, mask=rgba_preview.getchannel("A"))
        preview = background
    elif preview.mode == "L":
        preview = preview.convert("RGB")

    preview.thumbnail(CHAT_ASSET_PREVIEW_MAX_SIZE)
    buffer = BytesIO()
    preview.save(buffer, format="JPEG", quality=CHAT_ASSET_PREVIEW_QUALITY, optimize=True)
    return buffer.getvalue()
